import streamlit as st
import pandas as pd
import json
from datetime import datetime, date, timedelta
from openai import OpenAI
from jinja2 import Environment, FileSystemLoader, select_autoescape
import base64
from io import BytesIO
import os

# --- 导入 Word 导出库 ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 尝试导入 pdfminer.six，如果失败，在审计部分给出提示
try:
    from pdfminer.high_level import extract_text_to_fp
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False


# --- 全局 CSS 优化：解决 PDF 打印布局和文字大小问题 (已优化) ---
# 注入 CSS 以隐藏 Streamlit 默认 UI 元素，特别是打印时
st.markdown("""
<style>
/* 1. 全局字体优化 (保持网页显示不变) */
body {
    font-size: 11pt; 
}

/* 2. 打印模式优化 (@media print) */
@media print {
    /* 隐藏 Streamlit 的侧边栏、header、footer */
    .st-emotion-cache-vk3ypv, .st-emotion-cache-6crd03, .st-emotion-cache-12fmwpl, .st-emotion-cache-1wmy064, 
    [data-testid="stSidebar"], [data-testid="stHeader"], [data-testid="stToolbar"], footer, header { 
        display: none !important; 
        visibility: hidden !important;
    }
    
    /* 确保主要内容区最大化，移除 Streamlit 默认的 padding */
    .main, [data-testid="stAppViewBlockContainer"] { 
        padding-top: 0 !important; 
        padding-left: 15mm !important; /* 模拟 A4 边距 */
        padding-right: 15mm !important; /* 模拟 A4 边距 */
        padding-bottom: 0 !important;
        margin: 0 auto !important;
        max-width: 100% !important; 
    }
    
    /* 优化报告内 Markdown 文本的字体和行高 */
    h1, h2, h3, p, li {
        font-family: 'SimSun', 'Microsoft YaHei', sans-serif; /* 统一字体 */
        font-size: 11pt !important; /* 提高到 11pt 以适应 A4 纸的阅读大小 */
        line-height: 1.6; /* 增加行高以改善阅读体验 */
    }
    
    /* 保持标题大小相对一致 */
    h1 { font-size: 20pt !important; } /* 略微增大主标题 */
    h2 { font-size: 16pt !important; border-bottom: 1px solid #000; } /* 增大二级标题 */
    h3 { font-size: 14pt !important; color: #333; } /* 增大三级标题 */
    
    /* 确保表格不会被切断 */
    table {
        page-break-inside: avoid;
    }
}
</style>
""", unsafe_allow_html=True)


# --- 翻译辅助函数 ---
def translate_report(text_data, target_lang, api_key):
    """使用 DeepSeek API 翻译报告文本"""
    if not api_key:
        return None, "API Key 缺失"
        
    try:
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        
        # 优化目标语言的名称，以便 AI 理解
        if target_lang == "中文 (默认)":
            return text_data, None
        elif target_lang == "English (英文)":
            lang = "English"
        elif target_lang == "日本語 (日文)":
            lang = "Japanese"
        else:
            lang = target_lang

        prompt = f"""
        你是一位专业的质量管理翻译员。请将以下 8D 报告中的核心内容准确地翻译成 {lang}。
        请保留原有的Markdown、列表和分段格式。
        
        **请务必保留文本中的分隔符 `***AI_EVAL_SEP***`，不要对其进行翻译或移除。**
        
        仅返回翻译后的文本，不要添加任何解释或额外的Markdown标记。
        
        内容:
        {text_data}
        """
        
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        
        return response.choices[0].message.content, None
        
    except Exception as e:
        return None, f"翻译调用出错: {e}"


# --- Word 导出辅助函数 (使用 python-docx) ---
def create_word_document(extracted_data, evaluation_markdown, file_stream, is_translated=False):
    """
    根据 AI 审计结果创建 Word (DOCX) 文档
    如果 is_translated=True，则 assumed extracted_data 和 evaluation_markdown 已经是翻译后的 Markdown 纯文本。
    """
    document = Document()
    
    # 设置基础样式
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11) # 提高字号，与打印优化保持一致

    document.add_heading('AI 审计后的 8D 报告', level=1)
    document.add_paragraph(f"审计日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # --- 1. 结构化 8D 报告预览 ---
    document.add_heading('1. 结构化 8D 报告预览', level=2)
    
    if is_translated:
        # 如果是翻译后的内容，直接将翻译后的结构化数据 Markdown 文本添加到文档
        # 注意: 如果分割失败，export_eval_for_word 会是空字符串
        content_lines = (extracted_data + "\n\n" + evaluation_markdown).split('\n')
        
        # 针对翻译失败（内容合并）的情况进行处理
        if evaluation_markdown == "":
             document.add_paragraph('--- 结构化数据与评估合并显示 ---')

        for line in content_lines:
            if line.startswith('## '):
                document.add_heading(line.replace('##', '').strip(), level=3)
            elif line.startswith('*'):
                document.add_paragraph(line.replace('*', '').strip(), style='List Bullet')
            elif line.strip():
                document.add_paragraph(line)
        
    else:
        # 如果是中文原始数据 (JSON)，使用表格和列表结构化展示
        # 添加核心数据表格
        table = document.add_table(rows=5, cols=2, style='Table Grid')
        table.cell(0, 0).text = '阶段'
        table.cell(0, 1).text = '内容'
        
        core_rows = [
            ('D1 (组长)', extracted_data.get("D1_TeamLeader", "N/A")),
            ('D2 (问题描述)', extracted_data.get("D2_Problem", "N/A")),
            ('D4 (根本原因)', extracted_data.get("D4_RootCause", "N/A")),
            ('D8 (结论)', extracted_data.get("D8_Conclusion", "N/A")),
        ]
        
        for i, (stage, content) in enumerate(core_rows):
            table.cell(i + 1, 0).text = stage
            table.cell(i + 1, 1).text = content
            
        # 添加行动项
        document.add_paragraph('\n临时围堵措施 (D3 ICA):', style='List Bullet')
        for item in extracted_data.get("D3_ICA", ["N/A"]):
            document.add_paragraph(item.get("action", "N/A"), style='List Bullet')

        document.add_paragraph('\n永久对策 (D5 PCA):', style='List Bullet')
        for item in extracted_data.get("D5_Actions", ["N/A"]):
            document.add_paragraph(item.get("action", "N/A"), style='List Bullet')

        # 添加 D6/D7
        document.add_paragraph('\n')
        document.add_paragraph(f"D6 验证结果: {extracted_data.get('D6_Verification', 'N/A')}")
        document.add_paragraph(f"D7 标准化: {extracted_data.get('D7_Standardization', 'N/A')}")
    
        # --- 2. AI 审计评价 ---
        document.add_heading('2. AI 审计评价', level=2)
        
        # 简单解析 Markdown 评价，并添加到 Word
        for line in evaluation_markdown.split('\n'):
            if line.startswith('###'):
                document.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                document.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('*'):
                document.add_paragraph(line.replace('*', '').strip(), style='List Bullet')
            elif line.strip():
                document.add_paragraph(line)

    document.save(file_stream)
    file_stream.seek(0)


# --- Jinja2 环境设置 (用于生成HTML报告) ---
REPORT_TEMPLATE_HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>8D 报告 - {{ data.d0.title }}</title>
    <style>
        /* 保持 D8 报告模板的 A4 打印样式 */
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 20px; font-size: 11pt; } /* 提高字号 */
        .container { max-width: 900px; margin: auto; border: 1px solid #ccc; padding: 15px; box-shadow: 2px 2px 8px #eee; }
        h1 { color: #0056b3; }
        h2 { border-bottom: 2px solid #0056b3; padding-bottom: 5px; color: #0056b3; margin-top: 20px; }
        table { width: 100%; border-collapse: collapse; margin-bottom: 15px; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; font-weight: bold; }
        .section-table td:nth-child(1) { width: 30%; background-color: #f9f9f9; font-weight: bold; }
        .status-Completed { background-color: #d4edda; color: #155724; } 
        .status-Overdue { background-color: #f8d7da; color: #721c24; font-weight: bold; } 
        .status-DueSoon { background-color: #fff3cd; color: #856404; } 
        .status-Open { background-color: #f0f0f0; } 
        
        @media print {
            .container {
                max-width: 100%;
                border: none;
                padding: 0;
                box-shadow: none;
                margin: 0;
            }
            @page {
                size: A4; 
                margin: 20mm; 
            }
            h2 {
                page-break-before: auto;
                page-break-after: avoid;
            }
            table {
                page-break-inside: avoid;
            }
            body { font-size: 11pt; } /* 确保打印体也是 11pt */
        }
    </style>
</head>
<body>
    <div class="container">
        <h1 style="text-align: center;">8D 解决问题报告</h1>
        <p style="text-align: center; border-bottom: 1px dashed #ccc; padding-bottom: 10px;">
            **项目**: {{ data.d0.title }} | **客户**: {{ data.d0.customer }} | **日期**: {{ today }}
        </p>

        <h2>D1 & D2: 团队与问题描述</h2>
        <table class="section-table">
            <tr><td>**项目标题**</td><td>{{ data.d0.title }}</td></tr>
            <tr><td>**组长 (D1)**</td><td>{{ data.d1.leader }}</td></tr>
            <tr><td>**问题 (What)**</td><td>{{ data.d2.what }}</td></tr>
            <tr><td>**发生地点 (Where)**</td><td>{{ data.d2.where }}</td></tr>
            <tr><td>**详细描述**</td><td>{{ data.d2.desc }}</td></tr>
        </table>

        <h2>D3: 临时围堵措施 (ICA)</h2>
        {% if data.d3 %}
        <table>
            <tr><th>#</th><th>措施内容</th></tr>
            {% for item in data.d3 %}
            <tr><td>{{ loop.index }}</td><td>{{ item }}</td></tr>
            {% endfor %}
        </table>
        {% else %}
        <p>未录入临时围堵措施。</p>
        {% endif %}

        <h2>D4: 根本原因 (RCA)</h2>
        <table class="section-table">
            {% for i in range(5) %}
            <tr><td>**Why {{ i+1 }}**</td><td>{{ data.d4.whys[i] if data.d4.whys[i] else 'N/A' }}</td></tr>
            {% endfor %}
            <tr><td>**根本原因总结**</td><td>{{ data.d4.root_cause }}</td></tr>
        </table>

        <h2>D5/D6: 永久对策与实施</h2>
        {% if permanent_actions|length > 0 %}
        <table>
            <tr><th>对策内容</th><th>实施日期</th><th>状态</th></tr>
            {% for action in permanent_actions %}
            <tr class="status-{{ action.status_class }}">
                <td>{{ action.action }}</td>
                <td>{{ action.date }}</td>
                <td>{{ action.status_display }}</td>
            </tr>
            {% endfor %}
        </table>
        {% else %}
        <p>未录入永久对策。</p>
        {% endif %}

        <h2>D7 & D8: 预防与总结</h2>
        <table class="section-table">
            <tr><td>**FMEA/SOP更新 (D7)**</td><td>FMEA: {{ '✅' if data.d7.fmea else '❌' }} | CP: {{ '✅' if data.d7.cp else '❌' }} | SOP: {{ '✅' if data.d7.sop else '❌' }}</td></tr>
            <tr><td>**团队祝贺 (D8)**</td><td>8D 报告已完成并关闭，感谢团队的努力！</td></tr>
        </table>

    </div>
</body>
</html>
"""
env = Environment(loader=FileSystemLoader("."), autoescape=select_autoescape(['html', 'xml']))


# --- 页面基本配置 ---
st.set_page_config(page_title="8D 报告智能辅助系统 (DeepSeek)", layout="wide")

# --- Session State 初始化 (数据仓库) ---
if 'data' not in st.session_state:
    st.session_state.data = {
        'd0': {'title': '', 'customer': ''}, 
        'd1': {'leader': '', 'members': ''}, 
        'd2': {'what': '', 'where': '', 'desc': ''}, 
        'd3': [], 
        'd4': {'whys': ['', '', '', '', ''], 'root_cause': '', 'ai_analysis': None}, 
        'd5': [], 
        'd7': {'fmea': False, 'cp': False, 'sop': False}, 
        'd8': {}
    }
if 'audit_result' not in st.session_state:
    st.session_state.audit_result = {'extracted_data': None, 'evaluation_markdown': None, 'translated_data': None, 'translated_eval': None}


# --- 辅助函数：标题 ---
def section(title):
    st.markdown(f"## {title}")
    st.markdown("---")

# --- 辅助函数：状态逻辑处理 (用于 D5 条件格式) ---
def get_action_status(action_date_str, current_status):
    """根据日期字符串和当前状态判断最终状态并返回CSS类名和显示文本"""
    if current_status == 'Completed':
        return "Completed", "已完成"
        
    try:
        action_date = datetime.strptime(action_date_str, '%Y-%m-%d').date()
        today = date.today()
        
        if action_date < today:
            return "Overdue", "逾期/待验证"
        elif action_date <= today + timedelta(days=7):
            return "DueSoon", "临期"
        else:
            return "Open", "进行中"
    except:
        return "Open", "日期未设置"


# --- 侧边栏：导航与配置 (优化后的代码) ---
with st.sidebar:
    st.title("🚀 8D 流程智能辅助系统")
    
    with st.expander("🔑 AI 配置 (DeepSeek)", expanded=True):
        st.session_state['api_key'] = st.text_input("DeepSeek API Key", 
                                                    type="password", 
                                                    help="请填入 sk-开头的密钥", 
                                                    value=st.session_state.get('api_key', ''))
        base_url = "https://api.deepseek.com" 
    
    st.markdown("---")
    
    # --- 核心功能选择 (使用 Selectbox 进行分组) ---
    main_function = st.selectbox("选择核心功能", 
        ["1. 📝 新建/编辑 8D 报告", "2. 🔎 智能审计外部报告"], 
        index=0) 
    
# ----------------------------------------------


# ================= 业务逻辑 =================

if main_function == "1. 📝 新建/编辑 8D 报告":
    st.subheader("🛠️ 8D 报告创建流程")
    
    # --- 顶部导航栏 (使用 st.tabs 替代 st.radio) ---
    tab_names = ["D0: 准备", "D1: 团队", "D2: 问题描述", "D3: 围堵", 
                 "D4: 根本原因(AI)", "D5/D6: 对策", "D7: 预防", "D8: 报告生成"]
    tabs = st.tabs(tab_names)
    
    # 使用索引来判断当前所在的 Tab
    if st.session_state.get('current_tab_index', 0) >= len(tab_names):
        st.session_state['current_tab_index'] = 0 # 防止索引越界
        
    for i, tab in enumerate(tabs):
        with tab:
            step = tab_names[i]
            
            if step == "D0: 准备":
                section("D0: 基础信息")
                c1, c2 = st.columns(2)
                
                input_title = c1.text_input("报告标题", value=st.session_state.data['d0'].get('title', ''))
                input_customer = c2.text_input("客户名称", value=st.session_state.data['d0'].get('customer', ''))
                
                st.session_state.data['d0']['title'] = input_title
                st.session_state.data['d0']['customer'] = input_customer

            elif step == "D1: 团队":
                section("D1: 成立小组")
                
                input_leader = st.text_input("组长 (Leader)", value=st.session_state.data['d1'].get('leader', ''))
                input_members = st.text_area("成员名单", value=st.session_state.data['d1'].get('members', ''))
                
                st.session_state.data['d1']['leader'] = input_leader
                st.session_state.data['d1']['members'] = input_members

            elif step == "D2: 问题描述":
                section("D2: 问题描述 (AI分析的基础)")
                st.info("💡 提示：这里写得越详细，AI 分析得越准！")
                
                c1, c2 = st.columns(2)
                
                input_what = c1.text_input("发生了什么 (What)", value=st.session_state.data['d2'].get('what', ''))
                input_where = c2.text_input("发生在哪里 (Where)", value=st.session_state.data['d2'].get('where', ''))
                input_desc = st.text_area("详细描述整个过程", height=100, value=st.session_state.data['d2'].get('desc', ''))
                
                st.session_state.data['d2']['what'] = input_what
                st.session_state.data['d2']['where'] = input_where
                st.session_state.data['d2']['desc'] = input_desc

            elif step == "D3: 围堵":
                section("D3: 临时措施")
                
                new_ica = st.text_input("添加一条临时措施", key="new_ica_input")
                if st.button("➕ 添加") and new_ica:
                    st.session_state.data['d3'].append(new_ica)
                    st.rerun() 
                    
                if st.session_state.data['d3']:
                    st.write("📋 **已添加措施列表：**")
                    df = pd.DataFrame({'措施内容': st.session_state.data['d3']})
                    st.table(df)
                    
                    if st.button("🗑️ 清空列表"):
                        st.session_state.data['d3'] = []
                        st.rerun()

            elif step == "D4: 根本原因(AI)":
                section("D4: 根本原因分析 (DeepSeek 驱动)")
                
                d2_info = st.session_state.data['d2']
                problem_text = f"问题：{d2_info.get('what', '')}, 详情：{d2_info.get('desc', '')}"
                api_key = st.session_state.get('api_key')
                
                if not d2_info.get('what'):
                    st.warning("⚠️ 请先去 D2 步骤填写问题描述！")
                else:
                    # === AI 按钮 ===
                    if st.button("🤖 呼叫 DeepSeek 帮我分析", type="primary"):
                        if not api_key:
                            st.error("请在左侧边栏填入你的 DeepSeek API Key")
                        else:
                            try:
                                with st.spinner("DeepSeek 正在思考中..."):
                                    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
                                    
                                    prompt = f"""
                                    你是一个质量管理专家。请根据以下问题："{problem_text}"
                                    按 JSON 格式输出分析结果，包含两个字段：
                                    1. "five_whys": 一个包含5个字符串的列表，代表5个为什么的路径。
                                    2. "root_cause": 一句话总结根本原因。
                                    不要输出多余的 Markdown 标记。
                                    """
                                    
                                    response = client.chat.completions.create(
                                        model="deepseek-chat",
                                        messages=[{"role": "user", "content": prompt}],
                                        response_format={"type": "json_object"}
                                    )
                                    
                                    result = json.loads(response.choices[0].message.content)
                                    st.session_state.data['d4']['ai_analysis'] = result
                                    
                            except Exception as e:
                                st.error(f"AI 调用出错啦: {e}")

                    # === 显示结果与一键采纳 ===
                    if st.session_state.data['d4'].get('ai_analysis'):
                        ai_res = st.session_state.data['d4']['ai_analysis']
                        
                        st.success("分析完成！")
                        st.write("AI 建议的 5 Whys 路径：")
                        for i, w in enumerate(ai_res['five_whys']):
                            st.info(f"{i+1}. {w}")
                        
                        if st.button("⚡ 觉得不错，一键填入下方表格"):
                            for i in range(5):
                                if i < len(ai_res['five_whys']):
                                    st.session_state.data['d4']['whys'][i] = ai_res['five_whys'][i]
                            st.session_state.data['d4']['root_cause'] = ai_res['root_cause']
                            st.session_state.data['d4']['ai_analysis'] = None # 清除显示
                            st.rerun() 

                st.markdown("---")
                st.write("📝 **正式 5 Whys 记录表**")
                
                for i in range(5):
                    val = st.session_state.data['d4']['whys'][i]
                    new_val = st.text_input(f"Why {i+1}", value=val)
                    st.session_state.data['d4']['whys'][i] = new_val
                
                saved_root = st.session_state.data['d4'].get('root_cause', '')
                new_root = st.text_area("根本原因总结", value=saved_root)
                st.session_state.data['d4']['root_cause'] = new_root

            elif step == "D5/D6: 对策":
                section("D5/D6: 永久对策与实施")
                
                c1, c2 = st.columns([3, 1])
                new_pca = c1.text_input("新增永久对策", key="new_pca_input")
                action_date = c2.date_input("计划实施日期", value=date.today() + timedelta(days=14))
                
                if st.button("➕ 添加 PCA") and new_pca:
                    st.session_state.data['d5'].append({
                        "action": new_pca, 
                        "date": action_date.strftime('%Y-%m-%d'),
                        "status": "Open" 
                    })
                    st.rerun()
                    
                if st.session_state.data['d5']:
                    st.markdown("### 措施列表 (点击复选框标记完成)")
                    
                    updated_d5 = []
                    for i, action in enumerate(st.session_state.data['d5']):
                        col1, col2, col3 = st.columns([0.1, 4, 1])
                        
                        is_completed = col1.checkbox("", value=action.get('status') == 'Completed', key=f"d5_chk_{i}")
                        
                        # 更新状态
                        if is_completed:
                             action['status'] = 'Completed'
                        elif action.get('status') == 'Completed' and not is_completed:
                             action['status'] = 'Open'

                        # 显示内容和日期
                        col2.markdown(f"**{action['action']}**")
                        # 重新计算状态显示
                        status_class, status_display = get_action_status(action.get('date', ''), action.get('status', 'Open'))
                        col3.markdown(f"**{status_display}** ({action['date']})")
                        
                        updated_d5.append(action)
                    
                    st.session_state.data['d5'] = updated_d5


            elif step == "D7: 预防":
                section("D7: 预防再发生")
                
                st.session_state.data['d7']['fmea'] = st.checkbox("更新 FMEA (失效模式分析)", value=st.session_state.data['d7'].get('fmea', False))
                st.session_state.data['d7']['cp'] = st.checkbox("更新 Control Plan (控制计划)", value=st.session_state.data['d7'].get('cp', False))
                st.session_state.data['d7']['sop'] = st.checkbox("更新 SOP (作业指导书)", value=st.session_state.data['d7'].get('sop', False))


            elif step == "D8: 报告生成":
                section("D8: 报告预览与导出")
                st.info("💡 报告已按专业格式排版，并包含行动项的条件格式。")
                
                # --- 翻译和格式化选项 ---
                st.subheader("🌐 报告导出选项")
                c1, c2 = st.columns(2)
                
                export_format = c1.selectbox("选择导出格式", ["HTML (方便预览)", "PDF (推荐：通过浏览器打印)", "Word (基础复制粘贴)"])
                
                # 语言选择
                translation_lang = c2.selectbox("翻译为可选语言", ["中文 (默认)", "English (英文)", "日本語 (日文)"])
                
                # 报告核心数据准备 (与之前一致)
                d = st.session_state.data
                
                # 1. 处理 D5/D6 数据，添加状态和CSS类
                permanent_actions_processed = []
                for action in d['d5']:
                    status_class, status_display = get_action_status(action.get('date', ''), action.get('status', 'Open'))
                    
                    permanent_actions_processed.append({
                        'action': action['action'],
                        'date': action.get('date', 'N/A'),
                        'status': action.get('status'), # 保持原始状态
                        'status_class': status_class,
                        'status_display': status_display
                    })

                # 2. 准备 Jinja2 模板数据
                template_data = {
                    'data': d,
                    'today': date.today().strftime('%Y-%m-%d'),
                    'permanent_actions': permanent_actions_processed
                }

                # 3. 渲染 HTML (原始版本)
                template = env.from_string(REPORT_TEMPLATE_HTML)
                html_output = template.render(template_data)
                
                final_html_to_export = html_output
                
                # --- 翻译逻辑 ---
                if translation_lang != "中文 (默认)":
                    if not st.session_state.get('api_key'):
                         st.error("请在侧边栏填入 DeepSeek API Key 以启用翻译功能。")
                    else:
                        # 提取 Markdown 格式的报告核心内容
                        markdown_actions = '\n'.join([f'- {act["action"]} (Due: {act["date"]}, Status: {act["status_display"]})' for act in permanent_actions_processed])
                        markdown_content = f"""
                        # 8D Report: {d["d0"]["title"]}
                        
                        ## D1 & D2: Team and Problem Description
                        - Leader: {d["d1"]["leader"]}
                        - Problem (What): {d["d2"]["what"]}
                        - Detailed Description: {d["d2"]["desc"]}

                        ## D3: Interim Containment Action (ICA)
                        {'- ' + '  \n- '.join(d['d3']) if d['d3'] else 'N/A'}

                        ## D4: Root Cause Analysis (RCA)
                        - Root Cause Summary: {d["d4"]["root_cause"]}

                        ## D5/D6: Permanent Corrective Actions (PCA) & Verification
                        {markdown_actions if permanent_actions_processed else 'N/A'}
                        
                        ## D7 & D8: Prevention and Conclusion
                        - Standardization Check: FMEA: {'✅' if d['d7']['fmea'] else '❌'} | CP: {'✅' if d['d7']['cp'] else '❌'} | SOP: {'✅' if d['d7']['sop'] else '❌'}
                        - Conclusion: Report Closed.
                        """
                        
                        with st.spinner(f"正在将报告翻译为 {translation_lang}..."):
                            translated_markdown, error = translate_report(markdown_content, translation_lang, st.session_state.get('api_key'))
                            
                            if error:
                                st.error(f"翻译失败: {error}")
                            else:
                                st.subheader(f"✅ 翻译后的 Markdown 报告 (预览 - {translation_lang})")
                                st.markdown(translated_markdown)
                                st.markdown("---")
                                
                                # 替换最终导出的 HTML 为翻译后的 HTML 内容
                                final_html_to_export = f"""
                                <html><head><style>
                                body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 20px; font-size: 11pt; }}
                                /* 从原模板复制 A4 样式 */
                                {REPORT_TEMPLATE_HTML.split('<style>')[1].split('</style>')[0]}
                                </style></head>
                                <body><div class="container">{translated_markdown}</div></body></html>
                                """


                # 5. 提供下载链接 (根据选择的格式)
                
                # HTML 下载
                html_b64 = base64.b64encode(final_html_to_export.encode('utf-8')).decode()
                file_name_tag = translation_lang.split(' ')[0]
                file_name = f"8D_Report_{d['d0']['title']}_{file_name_tag}.html"
                href = f'<a href="data:text/html;charset=utf-8;base64,{html_b64}" download="{file_name}">📥 下载 {translation_lang} 版本的 HTML 报告</a>'
                st.markdown(href, unsafe_allow_html=True)
                
                st.markdown("---")
                st.subheader("导出操作指南：")
                
                if export_format == "PDF (推荐：通过浏览器打印)":
                    st.success("🎉 **已启用 A4 布局和 11pt 字体优化！**")
                    st.markdown("1. **下载 HTML 报告** (上面的链接)。")
                    st.markdown("2. **用浏览器打开** (Chrome/Edge/Firefox)。")
                    st.markdown("3. 按 **Ctrl+P (或 Cmd+P)** 打开打印对话框。")
                    st.markdown("4. 在目标打印机中选择 **“另存为 PDF”**。您将得到一张 A4 标准格式的专业报告。")
                
                elif export_format == "Word (基础复制粘贴)":
                    st.warning("Word 导出为手动过程。")
                    st.markdown("1. **下载 HTML 报告** (上面的链接)。")
                    st.markdown("2. **用浏览器打开**，复制所有内容。")
                    st.markdown("3. **粘贴到 Word 文档中**进行最后的格式调整。")
                
                else:
                    st.info("请选择导出格式查看具体操作指南。")


elif main_function == "2. 🔎 智能审计外部报告":
    
    st.subheader("🌟 8D 报告智能审计与评估")
    
    # 增加一个按钮/链接来触发打印，虽然本质上还是 Ctrl+P，但引导更明确
    st.info("💡 **布局优化完成：** 现在使用 Ctrl+P 打印时，侧边栏等冗余 UI 元素会被隐藏，布局将更清晰。")
    
    if st.session_state.audit_result.get('extracted_data'):
        # 只有在有报告内容时，才显示打印优化提示
        st.markdown('<button onclick="window.print()" style="font-size: 16px; padding: 10px 20px; background-color: #4CAF50; color: white; border: none; border-radius: 5px; cursor: pointer;">🖨️ 打印当前审计报告 (优化布局)</button>', unsafe_allow_html=True)
        st.markdown("---")
    
    
    st.info("💡 **功能说明：** 上传 TXT 或 PDF 文件，或直接粘贴报告文本。AI 将结构化解析并评价报告的完整性与逻辑。")
    if not PDF_SUPPORT:
        st.warning("⚠️ **缺少 PDF 支持库**：如需解析 PDF 文件，请先在命令行运行 `pip install pdfminer.six`")

    # 1. 文件上传区域
    uploaded_file = st.file_uploader("上传 8D 报告文件", type=['txt', 'pdf'])
    report_text = st.text_area("或者直接粘贴报告文本到这里", height=200, key="audit_text_input")
    
    # 文件处理逻辑 (保持不变)...
    if uploaded_file:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        if file_extension == 'pdf':
            if PDF_SUPPORT:
                try:
                    output_string = BytesIO()
                    uploaded_file.seek(0)
                    extract_text_to_fp(uploaded_file, output_string)
                    report_text = output_string.getvalue().decode('utf-8')
                    st.success("PDF 文件文本提取成功！")
                except Exception as e:
                    st.error(f"PDF 文本提取失败。错误：{e}")
                    report_text = None
            else:
                st.warning("请先安装 `pdfminer.six` 以支持 PDF 解析。")
                report_text = None
        
        elif file_extension == 'txt':
            try:
                uploaded_file.seek(0)
                report_text = uploaded_file.read().decode("utf-8")
                st.success("TXT 文件读取成功！")
            except Exception as e:
                st.error(f"TXT 文件读取失败: {e}")
                report_text = None
    
    # 审计执行逻辑
    if report_text and st.button("🚀 开始 AI 审计"):
        
        api_key = st.session_state.get('api_key')
        if not api_key:
             st.error("请在侧边栏输入 DeepSeek API Key。")
             st.stop()
        
        try:
            with st.spinner("DeepSeek 正在解析和审计报告..."):
                client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
                
                # --- 第一步：结构化提取 (优化 JSON 结构，要求 AI 提取行动项的 Owner 和 Due Date) ---
                extraction_prompt = f"""
                你是一个精准的 8D 信息提取机器人。请从以下 8D 报告文本中，提取出 D1-D8 的关键数据。
                
                请注意：
                1. D2 问题描述请分解为 5W2H 结构。
                2. D3 和 D5 的行动项，请尽量解析出 "action" (措施内容), "owner" (负责人), "dueDate" (计划完成日期，格式 YYYY-MM-DD 或 N/A), "status" (状态，Completed 或 Open)。如果无法解析 Owner/Date/Status，则使用 "N/A" 或 "Open"。
                3. D4 请同时提取 "OccurrenceRootCause" (发生根本原因) 和 "EscapeRootCause" (逃逸根本原因)。
                
                必须以合法的 JSON 格式输出，结构如下：
                {{
                    "D1_TeamLeader": "提取的组长姓名",
                    "D2_5W2H": {{
                        "What": "发生了什么",
                        "When": "何时发生",
                        "Where": "何处发生",
                        "Who": "谁发现/受影响",
                        "Why": "为什么发生 (初步原因)",
                        "How": "如何确认/测量",
                        "HowMuch": "影响范围/损失"
                    }},
                    "D3_ICA": [
                        {{"action": "临时措施1", "owner": "N/A", "dueDate": "N/A", "status": "Open"}},
                        // ... 更多 D3 措施
                    ],
                    "D4_RootCause": {{
                        "OccurrenceRootCause": "提取的发生根本原因",
                        "EscapeRootCause": "提取的逃逸根本原因"
                    }},
                    "D5_Actions": [
                        {{"action": "永久对策1", "owner": "N/A", "dueDate": "N/A", "status": "Open"}},
                        // ... 更多 D5 措施
                    ],
                    "D6_Verification": "提取的D6验证结果和数据总结",
                    "D7_Standardization": "提取的D7标准化措施（如FMEA/SOP更新）",
                    "D8_Conclusion": "提取的D8总结与团队祝贺"
                }}
                报告文本：{report_text}
                """
                
                extraction_response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "user", "content": extraction_prompt}],
                    response_format={"type": "json_object"},
                    temperature=0.1 
                )
                
                extracted_data = json.loads(extraction_response.choices[0].message.content)

                # --- 第二步：逻辑完整性与阶段评价 (全面审计) ---
                evaluation_prompt = f"""
                你是一位专业的 8D 流程审计师。请基于以下提取的 8D 数据，对报告的**完整性**和**逻辑性**进行阶段性评估。
                
                **提取数据：**
                {json.dumps(extracted_data, ensure_ascii=False, indent=2)}
                
                请用简洁的 Markdown 格式输出评估结果，并对**所有关键阶段**给出评价和改进建议：
                
                ## 8D 报告阶段性评价 (AI Audit)
                
                ### D0 & D1 (基础与团队)
                * **D0/D1 完整性**: 报告的基本信息（如标题、日期）和团队（组长、成员）是否明确记录？
                * **建议**: 
                
                ### D2 (问题描述)
                * **清晰度**: 问题描述 (D2_5W2H) 是否要素齐全？是否有量化的数据支持？
                * **建议**: 
                
                ### D3 (临时围堵措施 ICA)
                * **有效性**: 临时措施 (D3_ICA) 是否足够有力，能够彻底隔离所有不合格品，防止其流出？
                * **建议**: 
                
                ### D4 (根本原因 RCA)
                * **深度与逃逸点**: 根因是否区分了发生原因和逃逸原因？是否深入到体系或管理流程层面？
                * **建议**: 
                
                ### D5 & D6 (永久对策 PCA 与验证)
                * **逻辑关联**: **这是最重要的评估点。** 永久对策 (D5_Actions) 是否直接、彻底、一对一地消除了根本原因 (D4)？
                * **行动项管理**: D5 行动项中是否包含了 Owner、Due Date 或 Status 等信息？D6 验证是否明确、量化？
                * **建议**: 
                
                ### D7 & D8 (预防与总结)
                * **D7 标准化**: 是否明确提到了 FMEA/SOP/Control Plan 等文件的更新？这是确保流程预防的核心措施。
                * **D8 结论**: 报告关闭是否及时和得当？是否进行了团队祝贺？
                * **建议**: 
                """
                
                evaluation_response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "user", "content": evaluation_prompt}],
                    temperature=0.3
                )
                evaluation_markdown = evaluation_response.choices[0].message.content
                
                # 保存结果到 Session State，并清除旧的翻译结果
                st.session_state.audit_result['extracted_data'] = extracted_data
                st.session_state.audit_result['evaluation_markdown'] = evaluation_markdown
                st.session_state.audit_result['translated_data'] = None
                st.session_state.audit_result['translated_eval'] = None

        except Exception as e:
            st.error(f"AI 审计失败，请检查 API Key 或输入格式。错误信息: {e}")

    # --- 审计结果展示与导出 (已优化) ---
    
    extracted_data = st.session_state.audit_result.get('extracted_data')
    evaluation_markdown = st.session_state.audit_result.get('evaluation_markdown')
    
    if extracted_data and evaluation_markdown:
        
        st.markdown("---")
        
        # 翻译功能区 (保持不变)
        st.subheader("🌐 审计报告翻译")
        c_lang, c_btn = st.columns([1, 1])
        
        target_lang_text = c_lang.selectbox("选择审计报告的目标翻译语言", ["中文 (默认)", "English (英文)", "日本語 (日文)"], key="audit_lang_select")
        
        if c_btn.button("✨ 翻译审计报告", type="secondary"):
            api_key = st.session_state.get('api_key')
            if not api_key:
                c_btn.error("请在侧边栏输入 DeepSeek API Key。")
            else:
                # 1. 组合待翻译的原始文本 (Markdown 格式)
                # 注：为了让翻译模型更好地处理，此处不再进行复杂的 JSON 结构到 Markdown 的转换，而是直接将核心提取数据和评估Markdown合并
                # 目标是确保分隔符被保留
                data_markdown = f"""
                # 结构化 8D 报告核心内容
                ## D1/D2: {extracted_data.get("D1_TeamLeader", "N/A")} | {extracted_data.get("D2_5W2H", {}).get("What", "N/A")}
                ## D4 根本原因: 发生原因: {extracted_data.get("D4_RootCause", {}).get("OccurrenceRootCause", "N/A")} | 逃逸原因: {extracted_data.get("D4_RootCause", {}).get("EscapeRootCause", "N/A")}
                ## D8 结论: {extracted_data.get("D8_Conclusion", "N/A")}
                """
                full_content_to_translate = data_markdown + "\n\n***AI_EVAL_SEP***\n\n" + evaluation_markdown
                
                # 2. 调用翻译 API
                with st.spinner(f"正在将报告翻译为 {target_lang_text}..."):
                    translated_content, error = translate_report(full_content_to_translate, target_lang_text, api_key)
                    
                    if error:
                        st.error(f"翻译失败: {error}")
                        st.session_state.audit_result['translated_data'] = None
                        st.session_state.audit_result['translated_eval'] = None
                    else:
                        # 3. 分割翻译结果
                        parts = translated_content.split('\n\n***AI_EVAL_SEP***\n\n', 1)
                        if len(parts) == 2:
                            st.session_state.audit_result['translated_data'] = parts[0]
                            st.session_state.audit_result['translated_eval'] = parts[1]
                        else:
                            st.session_state.audit_result['translated_data'] = None 
                            st.session_state.audit_result['translated_eval'] = translated_content 
                            st.warning("⚠️ 翻译模型未保留结构化分隔符，结构化数据和评估已合并，请滚动查看下方完整内容。")
                        st.success(f"翻译完成，目标语言：{target_lang_text}")
        
        
        # --- 根据选择，确定最终显示和导出的内容 ---
        is_translated_content = (target_lang_text != "中文 (默认)") and st.session_state.audit_result.get('translated_eval') is not None
        
        if is_translated_content:
            
            if st.session_state.audit_result['translated_data']:
                # 成功分割
                st.subheader(f"📑 结构化 8D 报告预览 ({target_lang_text} 翻译结果)")
                st.markdown(st.session_state.audit_result['translated_data'])
                
                st.markdown("---")
                st.subheader(f"🧐 8D 报告阶段性评估 ({target_lang_text} 翻译结果)")
                st.markdown(st.session_state.audit_result['translated_eval'])
                
                export_data_for_word = st.session_state.audit_result['translated_data']
                export_eval_for_word = st.session_state.audit_result['translated_eval']
                
            else:
                # 分割失败，显示全部内容
                st.subheader(f"📑 结构化数据 & 评估合并报告 ({target_lang_text} 翻译结果)")
                st.markdown(st.session_state.audit_result['translated_eval'])
                
                export_data_for_word = st.session_state.audit_result['translated_eval']
                export_eval_for_word = "" 
            
        else:
            # --- 优化后的原始中文内容展示 ---
            st.subheader("📑 结构化 8D 报告预览 (原始中文)")
            
            # --- D1 (组长) & D8 (结论) ---
            st.markdown("### 👥 D1 团队 & D8 结论")
            d1_d8_data = {
                "阶段": ["D1 (组长)", "D8 (结论)"],
                "内容": [
                    extracted_data.get("D1_TeamLeader", "N/A"),
                    extracted_data.get("D8_Conclusion", "N/A")
                ]
            }
            st.table(pd.DataFrame(d1_d8_data).set_index('阶段'))
            
            # --- D2 (5W2H) ---
            st.markdown("### ❓ D2 问题描述 (5W2H)")
            d2_5w2h = extracted_data.get("D2_5W2H", {})
            d2_items = {
                "要素": ["What (何事)", "When (何时)", "Where (何处)", "Who (何人)", "Why (初步原因)", "How (如何确认)", "HowMuch (影响)"],
                "内容": [
                    d2_5w2h.get("What", "N/A"),
                    d2_5w2h.get("When", "N/A"),
                    d2_5w2h.get("Where", "N/A"),
                    d2_5w2h.get("Who", "N/A"),
                    d2_5w2h.get("Why", "N/A"),
                    d2_5w2h.get("How", "N/A"),
                    d2_5w2h.get("HowMuch", "N/A")
                ]
            }
            st.table(pd.DataFrame(d2_items).set_index('要素'))


            # --- D4 (根本原因 - 紧凑化展示) ---
            st.markdown("### 🔬 D4 根本原因 (发生与逃逸)")
            d4_root = extracted_data.get("D4_RootCause", {})
            st.markdown(f"**发生根本原因 (Occurrence):** {d4_root.get('OccurrenceRootCause', 'N/A')}")
            st.markdown(f"**逃逸根本原因 (Escape):** {d4_root.get('EscapeRootCause', 'N/A')}")
            
            # --- D3/D5/D6 行动项表格 ---
            st.markdown("### 🛠️ D3/D5/D6 行动项与验证")
            
            # 1. D3 表格
            st.markdown("##### D3 临时围堵措施 (ICA)")
            d3_actions = extracted_data.get("D3_ICA", [])
            if d3_actions and isinstance(d3_actions, list) and d3_actions[0].get("action"):
                df_d3 = pd.DataFrame(d3_actions)
                # 确保 Owner, DueDate, Status 存在，如果 AI 无法提取，则为 N/A 或 Open
                df_d3 = df_d3.rename(columns={'action': '措施内容', 'owner': '负责人', 'dueDate': '计划日期', 'status': '状态'})
                st.dataframe(df_d3)
            else:
                st.markdown("未提取到 D3 临时措施或格式不匹配。")


            # 2. D5 表格
            st.markdown("##### D5 永久对策 (PCA)")
            d5_actions = extracted_data.get("D5_Actions", [])
            if d5_actions and isinstance(d5_actions, list) and d5_actions[0].get("action"):
                df_d5 = pd.DataFrame(d5_actions)
                df_d5 = df_d5.rename(columns={'action': '对策内容', 'owner': '负责人', 'dueDate': '计划日期', 'status': '状态'})
                st.dataframe(df_d5)
            else:
                st.markdown("未提取到 D5 永久对策或格式不匹配。")
            
            # 3. D6/D7
            st.markdown("##### D6 验证结果")
            st.markdown(extracted_data.get('D6_Verification', 'N/A'))
            
            st.markdown("##### D7 标准化")
            st.markdown(extracted_data.get('D7_Standardization', 'N/A'))
            
            st.markdown("---")

            st.subheader("🧐 8D 报告阶段性评估 (原始中文)")
            st.markdown(evaluation_markdown)
            
            # 导出内容 (原始 JSON/Markdown)
            export_data_for_word = extracted_data
            export_eval_for_word = evaluation_markdown
        
        # --- 导出功能区 ---
        st.markdown("### 📥 导出结构化报告")
        
        c_word, c_pdf_btn = st.columns(2)
        
        # Word 导出按钮
        docx_io = BytesIO()
        # 传递 is_translated_content 标记和内容
        create_word_document(export_data_for_word, export_eval_for_word, docx_io, is_translated_content)

        file_tag = "Audit" if target_lang_text == "中文 (默认)" else f"Audit_{target_lang_text.split(' ')[0]}"
        
        c_word.download_button(
            label=f"下载 {target_lang_text} Word (.docx) 文件",
            data=docx_io.getvalue(),
            file_name=f"AI_{file_tag}_Report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
        # PDF 导出指南 
        c_pdf_btn.markdown(f"**下载 {target_lang_text} PDF 报告：**")
        c_pdf_btn.markdown("请使用浏览器打印功能 (**Ctrl+P/Cmd+P**)，选择**'另存为 PDF'**。**现在布局已优化。**")